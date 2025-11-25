import docx
import matplotlib.pyplot as plt

#открытие файла
doc = docx.Document('lion.docx')

#тут будет рапсоложен текст
text = []

#список для букв
letters = {}

#чтение текста
for paragraph in doc.paragraphs:
    text.append(paragraph.text)

#разбиение на абзацы
full_text = '\n'.join(text)

#меняем буквы с больших на малые
full_text = full_text.lower()

#удаление символов чтобы не мешали считать
punc = '/?!.,"«»[](){}-–:;—_1234567á890äâѵ°ȃ=ôęï―ç…όù*îö№üóûòêèàé’'
for i in range(0, len(punc)) :
    if punc[i] in full_text: 
        full_text = full_text.replace(punc[i], '')


#счетчик для слов
#задаем искомое слово
word = str(input('какое ваше слово: '))
#находим количество слов
kol = full_text.count(word)
#вычисляем процент встречи этих чисел
percent = ((kol / len(full_text))*100)

#цикл для нахождения букв в тексте с помощью слов
for i in full_text:
    #слова в буквы
    for j in i:
        #если есть буква, то счетчик +1
        if j in letters:
            letters[j] += 1
        #если буквы нет, то добавляет ее в наш словарь словарь
        else:
            letters.update({j:1})

#создаем оси для графика
keys = list(letters.keys())
values = list(letters.values())

#делпем график через bar потому что именно он сделает нам график нужного типа
plt.bar(keys, values)

#называем оси и сам график
plt.xlabel("буквы")
plt.ylabel("Количество")
plt.title("Гистограмма количества букв")

#отображаем графика
plt.show()
#создаем таблицу
doc.save('lion.docx')

dox = docx.Document()
table = dox.add_table(rows = 2, cols = 3)
table.style = 'Table Grid'
#заполняем таблицу
table.cell(0, 0).text = "Слово"
table.cell(0, 1).text = "Количество повторений"
table.cell(0, 2).text = "% от остальных слов"

table.cell(1, 0).text = str(word)
table.cell(1, 1).text = str(kol)
table.cell(1, 2).text = str(percent)
#сохраняем
dox.save('Word.docx')

def main():
    pass

    if __name__ == '__main__':
        main()